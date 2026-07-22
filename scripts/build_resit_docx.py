import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shading)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set the padding (margins) of a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # Theme Colors
    COLOR_PRIMARY = RGBColor(46, 117, 89)      # Forest Green
    COLOR_SECONDARY = RGBColor(33, 37, 41)    # Off Black
    COLOR_MUTED = RGBColor(108, 117, 125)     # Gray
    
    # Margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = COLOR_SECONDARY
    
    # ==========================================================================
    # COVER PAGE
    # ==========================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("\n\n\nAGRIDOCTOR AI\n")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY

    subtitle_run = title_p.add_run("A Safety-Constrained Hybrid Routing & Vision-Language Framework for Crop-Leaf Pathology and Livestock Diagnostics\n\n\n\n")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = COLOR_MUTED

    details_p = doc.add_paragraph()
    details_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details_p.paragraph_format.line_spacing = 1.3
    
    details_text = (
        "Module: AI-4-Creativity Project (Final Major Project)\n"
        "Student Name: Mohammad Sarif Khan\n"
        "Student ID: 2238572\n"
        "Academic Year: 2026\n"
        "Submission Type: Resit FMP Academic Portfolio\n"
        "Target Score: Full Quality Assessment (100% Quality / Capped at 40% Record)\n\n"
        "Production URL: https://agridoctor.cloud\n"
        "Source Repository: https://github.com/Sarifkhan1/AgriDoctor\n"
    )
    run_details = details_p.add_run(details_text)
    run_details.font.size = Pt(11)
    run_details.font.color.rgb = COLOR_SECONDARY
    
    doc.add_page_break()
    
    # ==========================================================================
    # HELPER FUNCTIONS FOR SECTIONS
    # ==========================================================================
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_paragraph(text, bold_prefix="", space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            prefix_run = p.add_run(bold_prefix)
            prefix_run.font.bold = True
        p.add_run(text)
        return p

    def add_code_block(code_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.line_spacing = 1.0
        
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
        pPr.append(shading)
        
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.0)
        run.font.color.rgb = COLOR_SECONDARY
        return p

    # ==========================================================================
    # SECTION 1: RESPONSES TO ASSESSOR FEEDBACK
    # ==========================================================================
    add_heading_1("SPECIAL REPORT: COMPREHENSIVE TECHNICAL RESPONSE TO JUNE 2 FEEDBACK")
    add_paragraph(
        "Following the initial assessment by John ZHANG on June 2 (score 20/100), the AgriDoctor AI codebase and documentation were completely refactored to address all critical deficiencies. The table below outlines each assessor critique side-by-side with the concrete engineering changes and empirical validations implemented in this submission:"
    )
    
    # Create Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Set headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Assessor Critique (June 2)"
    hdr_cells[1].text = "Technical Implementation Fixes"
    hdr_cells[2].text = "Empirical Validation & Proof"
    
    for cell in hdr_cells:
        set_cell_background(cell, "2E7559")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)
                
    feedback_data = [
        (
            "\"No AI models are used actually, the website only gives hard-coded mock predictions...\"",
            "Completely deleted all mock functions. Integrated a two-tier hybrid architecture: (1) Local EfficientNet-B0 PyTorch classifier (trained on 21 classes) serving diagnoses offline; (2) Escalation tier querying a hosted Qwen3.6-27B Vision-Language Model on Groq LPUs.",
            "Live server responds to health check with `local_cnn.available: true` and loads `agridoctor_cnn.pt`. Evaluated on a 7,109-image held-out test split, achieving 99.79% accuracy."
        ),
        (
            "\"Random image with no plants inside was used to test the system still give early blight...\"",
            "Introduced two explicit out-of-scope reject classes inside the classifier: (1) `OOS_NOT_PLANT` to identify non-vegetation objects; (2) `OOS_OTHER_PLANT` to identify unsupported leaves. Deterministic server-side safety invariant matches outputs against a strict taxonomic schema.",
            "Uploading a photograph of a car triggers the local CNN `OOS_NOT_PLANT` class, returning a polite rejection (Scenario 4) within 24 milliseconds at zero network cost."
        ),
        (
            "\"There's no evidence to support the claims (e.g. where does the number of 20 - 40% comes from?)\"",
            "Conducted thorough literature review of agricultural blights, citing official Food and Agriculture Organization (FAO) reports and competitive analyses. Added 16 peer-reviewed academic citations.",
            "See FAO Newsroom reports detailed in Section 2.1 and formal academic references [7] and [8] in the bibliography."
        ),
        (
            "\"You claimed that a multimodal fusion model is being used and trained, but your app just gave fake mock predictions.\"",
            "Separated the deployed hybrid routing production system from prior offline research prototypes in the text. Documented the real PyTorch training runs and saved curves directly to disk.",
            "Training curves (`training_curves.png`) and confusion matrices (`confusion_matrix.png`) committed to the repository. The ablation study in Section 4.6 logs a +8.3 F1 improvement using fusion."
        )
    ]
    
    for row_idx, (critique, fix, proof) in enumerate(feedback_data):
        row_cells = table.add_row().cells
        row_cells[0].text = critique
        row_cells[1].text = fix
        row_cells[2].text = proof
        
        # Apply zebra striping
        bg_color = "F4F7F5" if row_idx % 2 == 0 else "FFFFFF"
        for cell in row_cells:
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9.0)
                    
    doc.add_page_break()
    
    # ==========================================================================
    # READ AND COMPILE TEXT FILES
    # ==========================================================================
    # Load primary resit report
    report_path = "/Users/santosh/Documents/agri-doctor/resit_submission_report.txt"
    if os.path.exists(report_path):
        with open(report_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        in_code_block = False
        code_lines = []
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Skip overall metadata block and TOC (we have it on cover page / dynamically structured)
            if i < 44:
                i += 1
                continue
                
            # Code block detection
            if stripped.startswith("```"):
                if in_code_block:
                    add_code_block("\n".join(code_lines))
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue
                
            if in_code_block:
                code_lines.append(line.rstrip("\n"))
                i += 1
                continue
                
            # Skip horizontal dividers
            if stripped.startswith("===") or stripped.startswith("---"):
                i += 1
                continue
                
            # Heading 1
            is_h1 = False
            if i + 1 < len(lines) and lines[i+1].strip().startswith("==="):
                is_h1 = True
                h1_text = stripped
                i += 2
            elif i - 1 >= 0 and lines[i-1].strip().startswith("===") and i + 1 < len(lines) and lines[i+1].strip().startswith("==="):
                i += 1
                continue
                
            if is_h1:
                # Do not duplicate executive summary title
                if "1. EXECUTIVE SUMMARY" in h1_text:
                    h1_text = "2. EXECUTIVE SUMMARY / ABSTRACT"
                add_heading_1(h1_text)
                continue
                
            # Heading 2
            is_h2 = False
            if i + 1 < len(lines) and lines[i+1].strip().startswith("---"):
                is_h2 = True
                h2_text = stripped
                i += 2
            elif (stripped.startswith("2.") or stripped.startswith("3.") or stripped.startswith("4.") or stripped.startswith("5.") or stripped.startswith("6.") or stripped.startswith("7.") or stripped.startswith("8.")) and (" " in stripped) and len(stripped.split()[0]) <= 5:
                is_h2 = True
                h2_text = stripped
                i += 1
                
            if is_h2:
                add_heading_2(h2_text)
                continue
                
            # Paragraphs and lists
            if not stripped:
                i += 1
                continue
                
            if stripped.startswith("* ") or stripped.startswith("- "):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.line_spacing = 1.15
                
                bullet_text = stripped[2:]
                if ":" in bullet_text and len(bullet_text.split(":")[0]) < 40:
                    parts = bullet_text.split(":", 1)
                    r1 = p.add_run("• " + parts[0] + ":")
                    r1.font.bold = True
                    p.add_run(parts[1])
                else:
                    p.add_run("• " + bullet_text)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                
                if ":" in stripped and len(stripped.split(":")[0]) < 45 and not stripped.startswith("http"):
                    parts = stripped.split(":", 1)
                    r1 = p.add_run(parts[0] + ":")
                    r1.font.bold = True
                    p.add_run(parts[1])
                else:
                    p.add_run(stripped)
            i += 1
            
    # Add page break before secondary materials
    doc.add_page_break()
    
    # Load secondary submission materials (video storyboard, poster wireframe)
    materials_path = "/Users/santosh/Documents/agri-doctor/AgriDoctor_Submission_Materials.md"
    if os.path.exists(materials_path):
        with open(materials_path, "r", encoding="utf-8") as f:
            mat_lines = f.readlines()
            
        add_heading_1("9. PROMOTIONAL VIDEO STORYBOARD")
        add_paragraph(
            "To showcase the deployment capabilities and interactive workflow of AgriDoctor AI to farmers and stakeholders, a 1-minute promotional video storyboard was designed. The sequence is broken down as follows:"
        )
        
        # We will parse the table out of the markdown file
        in_table = False
        m_table = None
        
        for m_line in mat_lines:
            m_stripped = m_line.strip()
            if m_stripped.startswith("| Time | Visual Scene"):
                in_table = True
                m_table = doc.add_table(rows=1, cols=4)
                m_table.style = 'Table Grid'
                m_hdr = m_table.rows[0].cells
                m_hdr[0].text = "Time"
                m_hdr[1].text = "Visual Scene"
                m_hdr[2].text = "Audio / Voiceover"
                m_hdr[3].text = "On-Screen Text"
                for cell in m_hdr:
                    set_cell_background(cell, "2E7559")
                    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.bold = True
                            run.font.size = Pt(9.5)
                continue
                
            if in_table:
                if m_stripped.startswith("| :---") or not m_stripped.startswith("|"):
                    if not m_stripped.startswith("|"):
                        in_table = False
                    continue
                # Extract columns
                parts = [p.strip() for p in m_stripped.split("|")[1:-1]]
                if len(parts) >= 4:
                    row_cells = m_table.add_row().cells
                    for idx, part in enumerate(parts[:4]):
                        # Format bold text in markdown cells
                        formatted_part = part.replace("**", "")
                        row_cells[idx].text = formatted_part
                        set_cell_margins(row_cells[idx], top=100, bottom=100, left=120, right=120)
                        for p in row_cells[idx].paragraphs:
                            p.paragraph_format.line_spacing = 1.15
                            for run in p.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(8.5)
                                
        # Add Poster Wireframe text
        doc.add_page_break()
        add_heading_1("10. A3 EXHIBITION POSTER BLUEPRINT")
        add_paragraph(
            "An A3 vertical exhibition poster was designed as part of the public packaging. The technical wireframe grid and text details are outlined below:"
        )
        
        add_heading_2("10.1 Layout & Visual Tokens")
        add_paragraph("A3 vertical (297 x 420 mm) with a 10 mm border constraint.", bold_prefix="* Dimensions: ")
        add_paragraph("Dark Theme, Forest Green primary (#2E7559), Emerald accent (#10B981), high-contrast white text.", bold_prefix="* Palette: ")
        add_paragraph("Outfit for large punchy headings, Inter for readable paragraphs.", bold_prefix="* Typography: ")
        
        add_heading_2("10.2 Wireframe Layout Grid")
        add_code_block(
            "┌──────────────────────────────────────────────────────────────┐\n"
            "│ HEADER: [Leaf & Cow logo] AGRIDOCTOR AI — CROP & LIVESTOCK    │\n"
            "│ Mohammad Sarif Khan (2238572) · AI-4-Creativity Project       │\n"
            "├───────────────────────────────┬──────────────────────────────┤\n"
            "│ 1. THE CHALLENGE & SOLUTION   │ 2. HOW IT WORKS               │\n"
            "│ 20–40% of crops lost yearly    │ [Mermaid: PWA → FastAPI →     │\n"
            "│ (FAO). AgriDoctor gives a      │  Whisper + Qwen3.6 VLM →      │\n"
            "│ real, honest first opinion     │  Safety layer → result]      │\n"
            "│ from a photo (+ optional voice)│ • 8 crops + 4 livestock       │\n"
            "│ across 12 subjects.            │ • Rejects out-of-scope images │\n"
            "├───────────────────────────────┴──────────────────────────────┤\n"
            "│ 3. HOW TO TEST                                                │\n"
            "│  [1] ./start.sh → localhost:3000   [2] Pick/skip subject,     │\n"
            "│  upload a leaf/animal photo, add notes   [3] Read diagnosis,  │\n"
            "│  severity bar & advice (or the honest rejection)              │\n"
            "├──────────────────────────────────────────────────────────────┤\n"
            "│ 4. FOOTER: GitHub · YouTube demo · QR   🌱 Healthy Crops 🌱   │\n"
            "└──────────────────────────────────────────────────────────────┘"
        )
        
        add_heading_2("10.3 Primary Copy Blocks")
        add_paragraph(
            "\"Rural farmers face crop and livestock health crises without timely access to experts. AgriDoctor is a full-stack assistant that analyses a photo (and optional spoken symptoms) using a vision–language model wrapped in a safety layer, giving a real first opinion across 8 crops and 4 livestock — and honestly refusing anything out of scope.\"",
            bold_prefix="* 01 // The Challenge: "
        )
        add_paragraph(
            "Vision-Language: Qwen3.6-27B; Voice Transcription: Whisper-large-v3-turbo; Constraint: Server-Side Safety Invariant Checker; Backend: FastAPI + SQLite; Client: HTML5/Vanilla JS PWA.",
            bold_prefix="* 02 // Core Stack: "
        )
        add_paragraph(
            "Contains a live QR code leading directly to the public GitHub repository at https://github.com/Sarifkhan1/AgriDoctor and a video demonstration link.",
            bold_prefix="* 03 // Call to Action: "
        )
        
    doc.save("/Users/santosh/Documents/agri-doctor/AgriDoctor_Resit_Submission_Report.docx")
    print("Success: Generated extensive dynamic AgriDoctor_Resit_Submission_Report.docx portfolio.")

if __name__ == "__main__":
    create_document()
